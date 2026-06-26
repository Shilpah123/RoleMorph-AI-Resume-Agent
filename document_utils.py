import os
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def read_docx(filepath: str) -> str:
    """Read text content from a Word document."""
    doc = Document(filepath)
    text_content = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_content.append(paragraph.text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_content.append(cell.text)
    
    return '\n'.join(text_content)


def read_file_smart(filepath: str) -> str:
    """Read file content, automatically detecting format (txt, md, or docx)."""
    file_ext = os.path.splitext(filepath)[1].lower()
    
    if file_ext == '.docx':
        return read_docx(filepath)
    else:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()


def write_docx(filepath: str, content: str, title: Optional[str] = None):
    """Write content to a Word document with professional formatting."""
    doc = Document()
    
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    if title:
        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        doc.add_paragraph()
    
    lines = content.split('\n')
    current_section = None
    
    for line in lines:
        line = line.strip()
        
        if not line:
            doc.add_paragraph()
            continue
        
        if line.isupper() and len(line) < 50:
            para = doc.add_paragraph(line)
            para.runs[0].font.bold = True
            para.runs[0].font.size = Pt(14)
            current_section = line
        
        elif line.endswith(':') and len(line) < 100:
            para = doc.add_paragraph(line)
            para.runs[0].font.bold = True
            para.runs[0].font.size = Pt(12)
        
        elif line.startswith('•') or line.startswith('-'):
            para = doc.add_paragraph(line.lstrip('•-').strip(), style='List Bullet')
            para.runs[0].font.size = Pt(11)
        
        else:
            para = doc.add_paragraph(line)
            para.runs[0].font.size = Pt(11)
            
            if current_section and any(keyword in line.lower() for keyword in ['email:', 'phone:', 'linkedin:', 'location:']):
                para.runs[0].font.size = Pt(10)
    
    doc.save(filepath)


def save_file_smart(filepath: str, content: str, title: Optional[str] = None):
    """Save content to file, automatically detecting format from extension."""
    file_ext = os.path.splitext(filepath)[1].lower()
    
    os.makedirs(os.path.dirname(filepath) if os.path.dirname(filepath) else '.', exist_ok=True)
    
    if file_ext == '.docx':
        write_docx(filepath, content, title)
    else:
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)


def extract_text_from_upload(uploaded_file) -> str:
    """Extract text from Streamlit uploaded file (supports txt, md, docx)."""
    file_ext = os.path.splitext(uploaded_file.name)[1].lower()
    
    if file_ext == '.docx':
        doc = Document(uploaded_file)
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)
        
        return '\n'.join(text_content)
    else:
        return uploaded_file.read().decode('utf-8')

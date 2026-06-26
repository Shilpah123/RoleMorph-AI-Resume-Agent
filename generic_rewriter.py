"""
Generic resume rewriter - works with ANY resume format.
Intelligently analyzes and enhances content without hardcoded patterns.
Now powered by multi-agent architecture.
"""

import os
import re
from typing import Dict, List, Tuple
from docx import Document
from docx.shared import RGBColor
from agents.resume_orchestrator import ResumeOrchestrator


class GenericRewriter:
    """Generic rewriter that works with any resume format."""
    
    def __init__(self):
        self.action_verbs = {
            'led': 'Spearheaded', 'managed': 'Directed', 'created': 'Developed',
            'made': 'Established', 'did': 'Executed', 'worked': 'Collaborated',
            'helped': 'Facilitated', 'improved': 'Enhanced', 'changed': 'Transformed',
            'built': 'Architected', 'wrote': 'Authored', 'handled': 'Managed'
        }
        self.orchestrator = ResumeOrchestrator()
        self.agent_analysis = None  # Store latest agent analysis
    
    def extract_skills(self, job_desc: str) -> List[str]:
        """Extract skills and requirements dynamically from job description."""
        requirements = set()
        
        # Extract degree requirements
        degree_patterns = [
            r'(bachelor|master|phd|doctorate|msc|bsc|life sciences|science degree)',
            r'(degree in \w+)',
            r'(\w+ degree)'
        ]
        for pattern in degree_patterns:
            matches = re.findall(pattern, job_desc, re.IGNORECASE)
            for match in matches:
                if isinstance(match, tuple):
                    match = match[0]
                if len(match) > 3:
                    requirements.add(match.lower())
        
        # Extract experience requirements
        exp_patterns = [
            r'(\d+[\+\-]?\s*years?\s+(?:of\s+)?(?:experience|exp)(?:\s+in\s+[\w\s]+)?)',
            r'(experience (?:in|with|of) [\w\s]{5,40})',
            r'(previous [\w\s]+ experience)',
            r'(proven [\w\s]+ experience)'
        ]
        for pattern in exp_patterns:
            matches = re.findall(pattern, job_desc, re.IGNORECASE)
            for match in matches:
                if isinstance(match, tuple):
                    match = match[0]
                cleaned = match.strip().lower()
                if 5 < len(cleaned) < 100:
                    requirements.add(cleaned)
        
        # Extract specific skills/tools mentioned
        skill_indicators = [
            r'(?:experience with|knowledge of|proficiency in|skilled in|expertise in)\s+([\w\s]{5,40})',
            r'(?:ability to|capable of)\s+([\w\s]{5,40})',
            r'(?:strong|excellent|good)\s+([\w\s]{5,30})\s+skills?'
        ]
        for pattern in skill_indicators:
            matches = re.findall(pattern, job_desc, re.IGNORECASE)
            for match in matches:
                cleaned = match.strip().lower()
                if 5 < len(cleaned) < 60:
                    requirements.add(cleaned)
        
        # Extract domain-specific terms (nouns that appear multiple times)
        words = re.findall(r'\b[A-Z][a-z]{3,}\b', job_desc)
        word_freq = {}
        for word in words:
            word_lower = word.lower()
            if word_lower not in ['the', 'and', 'for', 'with', 'this', 'that', 'from', 'have', 'will']:
                word_freq[word_lower] = word_freq.get(word_lower, 0) + 1
        
        # Add frequently mentioned terms (likely important domain terms)
        for word, freq in word_freq.items():
            if freq >= 2 and len(word) > 4:
                requirements.add(word)
        
        return sorted(list(requirements))[:30]  # Limit to top 30
    
    def get_agent_analysis(self, resume_text: str, job_desc: str) -> Dict:
        """Get detailed agent analysis without modifying resume."""
        return self.orchestrator.analyze_and_optimize(resume_text, job_desc)
    
    def extract_keywords(self, job_desc: str) -> List[str]:
        """Extract keywords from job description."""
        keywords = set()
        job_lower = job_desc.lower()
        
        # Content/writing-focused skills (for content roles)
        content_skills = ['technical writing', 'content development', 'content editing',
                         'content publishing', 'editorial', 'proofreading', 'copywriting',
                         'user documentation', 'user manuals', 'sme interviews',
                         'user feedback', 'web content', 'email content', 'internal communications',
                         'style guide', 'content standards', 'editorial review', 'cms']
        
        for skill in content_skills:
            if skill in job_lower:
                keywords.add(skill.title())
        
        # Configuration Management / Engineering Process skills
        cm_skills = ['configuration management', 'change management', 'change control',
                    'product lifecycle', 'plm', 'product data management', 'pdm',
                    'eco', 'mco', 'engineering change order', 'manufacturing change order',
                    'engineering change', 'manufacturing change',
                    'bom', 'bill of materials', 'workflow development',
                    'process governance', 'configuration governance', 'data governance',
                    'aras', 'agile plm', 'windchill', 'teamcenter', 'change control board',
                    'release management', 'product configuration']
        
        for skill in cm_skills:
            if skill in job_lower:
                keywords.add(skill.upper() if len(skill) <= 4 else skill.title())
        
        # Technical skills
        tech_skills = ['python', 'html', 'css', 'xml', 'markdown', 'dita',
                      'github', 'git', 'version control', 'visio', 'lucidchart',
                      'confluence', 'jira', 'markup languages']
        
        for skill in tech_skills:
            if skill in job_lower:
                keywords.add(skill.upper() if len(skill) <= 4 else skill.title())
        
        # Soft/process skills
        soft_skills = ['communication', 'collaboration', 'stakeholder management',
                      'cross-functional', 'agile', 'problem solving', 'time management',
                      'process design', 'learning quickly', 'process improvement',
                      'workflow automation', 'process governance', 'change leadership',
                      'transformation', 'standardization']
        
        for skill in soft_skills:
            if skill in job_lower:
                keywords.add(skill.title())
        
        return sorted(list(keywords))
    
    def is_bullet_point(self, text: str) -> bool:
        """Detect if a line is a bullet point - works with any format."""
        text = text.strip()
        
        # Empty or too short
        if len(text) < 15:
            return False
        
        # Starts with bullet character
        if text[0] in ['•', '-', '*', '○', '·', '►', '▪']:
            return True
        
        # Starts with tab + bullet
        if text.startswith('\t') and len(text) > 20:
            return True
        
        # Starts with number (numbered list)
        if re.match(r'^\d+[\.\)]\s+', text):
            return True
        
        # Looks like a bullet (short sentence, starts with action verb)
        words = text.split()
        if len(words) >= 3 and len(words) <= 30:
            first_word = words[0].lower().rstrip('.,;:')
            # Check if starts with action verb
            if first_word in ['led', 'managed', 'created', 'developed', 'built',
                            'designed', 'implemented', 'established', 'improved',
                            'coordinated', 'directed', 'executed', 'facilitated',
                            'spearheaded', 'drove', 'delivered', 'achieved']:
                return True
        
        return False
    
    def is_section_header(self, text: str) -> Tuple[bool, str]:
        """Detect if line is a section header. Returns (is_header, section_type)."""
        text_lower = text.lower().strip()
        
        # Too long to be a header
        if len(text.split()) > 6:
            return False, None
        
        # Check for common section headers
        if 'summary' in text_lower or 'profile' in text_lower or 'objective' in text_lower:
            return True, 'summary'
        
        if 'experience' in text_lower or 'employment' in text_lower or 'work history' in text_lower:
            return True, 'experience'
        
        if 'education' in text_lower or 'skills' in text_lower or 'competenc' in text_lower:
            return True, 'other'
        
        if 'award' in text_lower or 'certification' in text_lower or 'achievement' in text_lower:
            return True, 'other'
        
        return False, None
    
    def create_new_summary(self, job_desc: str, keywords: List[str], original: str = "") -> str:
        """Create new professional summary based on job description."""
        # Extract years of experience - ONLY from original resume
        years = 15  # default fallback
        
        # Extract candidate's actual background from original resume
        candidate_background = "delivering results"  # default fallback
        
        if original:
            # Extract years
            years_match = re.search(r'(\d+)\+?\s*years', original.lower())
            if years_match:
                years = int(years_match.group(1))
            else:
                years_match = re.search(r'(\d+)\s*years?\s+(?:of\s+)?(?:experience|in)', original.lower())
                if years_match:
                    years = int(years_match.group(1))
            
            # Extract candidate's role/background from original summary
            # Look for phrases like "Technical Writer", "Engineer", "Manager", etc.
            orig_lower = original.lower()
            
            # Common role patterns
            if 'technical writ' in orig_lower or 'technical commun' in orig_lower:
                candidate_background = "technical communication and documentation"
            elif 'content' in orig_lower and ('develop' in orig_lower or 'writ' in orig_lower):
                candidate_background = "content development and technical writing"
            elif 'engineer' in orig_lower:
                candidate_background = "engineering and technical development"
            elif 'manager' in orig_lower or 'management' in orig_lower:
                candidate_background = "management and leadership"
            elif 'analyst' in orig_lower:
                candidate_background = "analysis and problem-solving"
            elif 'developer' in orig_lower or 'development' in orig_lower:
                candidate_background = "development and implementation"
            elif 'specialist' in orig_lower:
                candidate_background = "specialized technical work"
            # If we find specific domain mentions, use those
            if 'documentation' in orig_lower:
                if 'technical' not in candidate_background:
                    candidate_background = candidate_background.replace('delivering results', 'documentation and technical communication')
        
        # Build simple, reliable summary
        parts = []
        
        # Opening that uses candidate's actual background
        parts.append(f"Experienced professional with {years}+ years in {candidate_background}.")
        
        # Add relevant keywords if found (filter out generic soft skills)
        if keywords:
            # Filter out overly generic keywords
            filtered_keywords = [k for k in keywords if k.lower() not in ['communication', 'collaboration', 'visio', 'lucidchart']]
            
            if filtered_keywords:
                top_keywords = filtered_keywords[:6]
                if len(top_keywords) <= 3:
                    parts.append(f"Expertise in {', '.join(top_keywords)}.")
                else:
                    parts.append(f"Expertise in {', '.join(top_keywords[:3])}, {', '.join(top_keywords[3:])}.")
        
        # Generic strength statement
        parts.append("Proven ability to collaborate with cross-functional teams, deliver high-quality documentation, and drive process improvements.")
        
        # Combine and limit to 120 words
        summary = " ".join(parts)
        words = summary.split()
        if len(words) > 120:
            summary = " ".join(words[:120]) + "."
        
        return summary
    
    def enhance_bullet(self, bullet: str, keywords: List[str]) -> str:
        """Enhance bullet with natural content-focused language - NO keyword stuffing."""
        # Remove bullet marker
        bullet = bullet.strip()
        for marker in ['•', '\t•', '-', '*', '○', '·', '►', '▪', '\t']:
            if bullet.startswith(marker):
                bullet = bullet.replace(marker, '', 1).strip()
        
        # Remove numbered list markers
        bullet = re.sub(r'^\d+[\.\)]\s+', '', bullet)
        
        # Enhance verb if starts with weak verb
        words = bullet.split()
        if words:
            first_word_lower = words[0].lower().rstrip('.,;:')
            if first_word_lower in self.action_verbs:
                words[0] = self.action_verbs[first_word_lower]
                bullet = ' '.join(words)
        
        # Analyze bullet content to add natural, contextual enhancements
        bullet_lower = bullet.lower()
        enhanced = bullet.rstrip('.')
        
        # Add natural content-focused outcomes with VARIATION to avoid repetition
        # Use different phrases for similar concepts
        
        # Track what we've added to avoid repetition across bullets
        import random
        
        # For documentation/content work - VARY the language with more diversity
        if any(term in bullet_lower for term in ['document', 'content', 'manual', 'guide']):
            if 'sme' not in bullet_lower and 'interview' not in bullet_lower:
                variations = [
                    ", converting SME input into structured documentation",
                    ", gathering requirements from subject matter experts",
                    ", conducting stakeholder interviews to capture technical requirements",
                    ", translating technical discussions into clear user content",
                    ", working with engineering teams to validate technical accuracy",
                    ", refining SME feedback into documentation updates",
                    ", partnering with product teams to document complex features",
                    ", extracting technical knowledge from cross-functional experts"
                ]
                # Filter out already used phrases
                available = [v for v in variations if v not in getattr(self, 'used_phrases', [])]
                if not available:
                    available = variations  # Reset if all used
                chosen = random.choice(available)
                if not hasattr(self, 'used_phrases'):
                    self.used_phrases = []
                self.used_phrases.append(chosen)
                enhanced += chosen
            elif 'clarity' not in bullet_lower and 'quality' not in bullet_lower and 'edit' not in bullet_lower:
                variations = [
                    ", improving readability and user comprehension",
                    ", enhancing documentation quality and structure",
                    ", ensuring content aligns with style guides and standards",
                    ", refining content based on user feedback",
                    ", reviewing and editing documentation for clarity and consistency"
                ]
                enhanced += random.choice(variations)
        
        # For migration/transformation work - use CM language if CM role detected
        elif any(term in bullet_lower for term in ['migration', 'transform', 'moderniz']):
            # Check if this is a CM-focused transformation or content transformation
            if any(cm_term in bullet_lower for cm_term in ['arbortext', 'dita', 'xml', 'structure', 'architecture']):
                # CM/Engineering transformation language
                variations = [
                    ", redesigning structured product information models and improving system scalability",
                    ", standardizing information structures to improve consistency and reuse",
                    ", aligning content structure with system and product requirements",
                    ", improving data flow and operational efficiency across engineering systems"
                ]
                enhanced += random.choice(variations)
            elif 'publish' not in bullet_lower and 'cms' not in bullet_lower:
                # Content publishing language
                variations = [
                    ", establishing publishing workflows and CMS integration",
                    ", implementing content management systems and publishing standards",
                    ", streamlining content delivery through web and intranet platforms"
                ]
                enhanced += random.choice(variations)
        
        # For process/workflow improvements - adapt language based on context
        elif any(term in bullet_lower for term in ['process', 'workflow', 'automat']):
            # Check if this is engineering/data workflow or editorial workflow
            if any(eng_term in bullet_lower for eng_term in ['python', 'automation', 'script', 'data', 'system']):
                # Engineering/CM workflow language
                variations = [
                    ", streamlining structured data handling and reducing manual processing effort",
                    ", supporting workflow stability and cross-team coordination",
                    ", improving data flow and operational efficiency",
                    ", enabling scalable process automation across engineering teams"
                ]
                enhanced += random.choice(variations)
            elif 'editorial' not in bullet_lower:
                # Content/editorial workflow language
                variations = [
                    ", implementing editorial review processes and quality assurance",
                    ", establishing content review cycles and approval workflows",
                    ", standardizing documentation processes and quality checks"
                ]
                enhanced += random.choice(variations)
        
        # For team/leadership work - adapt based on context
        elif any(term in bullet_lower for term in ['team', 'manag', 'lead']):
            if 'mentor' not in bullet_lower:
                # Check if engineering/technical team or content team
                if any(eng_term in bullet_lower for eng_term in ['developer', 'technical', 'product', 'engineering']):
                    # Engineering team leadership language
                    variations = [
                        ", delivering structured technical deliverables supporting engineering processes",
                        ", ensuring alignment of technical information across systems",
                        ", implementing standardized workflows improving consistency and efficiency",
                        ", coordinating cross-functional engineering and documentation efforts"
                    ]
                    enhanced += random.choice(variations)
                else:
                    # Content team leadership language
                    variations = [
                        ", mentoring writers and maintaining documentation standards",
                        ", coaching team members on content best practices",
                        ", fostering quality standards and team collaboration"
                    ]
                    enhanced += random.choice(variations)
        
        # For collaboration work - adapt language based on context
        elif any(term in bullet_lower for term in ['collaborat', 'partner', 'work with']):
            if 'feedback' not in bullet_lower and 'requirement' not in bullet_lower:
                # Check if engineering collaboration or general collaboration
                if any(eng_term in bullet_lower for eng_term in ['engineering', 'product', 'technical', 'hardware', 'software']):
                    # Engineering collaboration language
                    variations = [
                        ", ensuring alignment of technical information across systems and teams",
                        ", supporting engineering-facing workflows and system requirements",
                        ", coordinating with engineering stakeholders to ensure data accuracy",
                        ", aligning structured information with product lifecycle needs"
                    ]
                    enhanced += random.choice(variations)
                else:
                    # General collaboration language
                    variations = [
                        ", interpreting stakeholder requirements and user feedback",
                        ", incorporating user feedback to refine and improve content",
                        ", translating business requirements into clear documentation"
                    ]
                    enhanced += random.choice(variations)
        
        # For validation/quality work - add explicit editing for clarity
        elif any(term in bullet_lower for term in ['validat', 'review', 'ensure']):
            if 'proofread' not in bullet_lower and 'edit' not in bullet_lower:
                variations = [
                    ", proofreading and editing for accuracy and clarity",
                    ", conducting editorial reviews to ensure quality",
                    ", reviewing content for consistency and technical accuracy",
                    ", editing documentation for clarity and consistency"
                ]
                enhanced += random.choice(variations)
        
        # For awards/recognition - keep as is
        elif any(term in bullet_lower for term in ['award', 'recogni', 'honor']):
            pass  # Don't add anything to awards
        
        # Default: add general content writing outcome with variation
        else:
            if not any(word in bullet_lower for word in ['improv', 'enhanc', 'ensur', 'deliver']):
                variations = [
                    ", ensuring accuracy and timely delivery",
                    ", maintaining quality standards and meeting deadlines",
                    ", delivering high-quality content on schedule"
                ]
                enhanced += random.choice(variations)
        
        return enhanced
    
    def run_workflow(self, input_path: str, output_path: str, job_description: str) -> Dict:
        """Run generic workflow that works with any resume - now with multi-agent intelligence."""
        
        print("\n" + "="*70)
        print("RoleMorph: AI Resume Optimization Agent System")
        print("="*70)
        
        # Load document first to get resume text
        print("\nStep 1: Loading document...")
        doc = Document(input_path)
        print(f"  ✓ Loaded {len(doc.paragraphs)} paragraphs")
        
        # Extract full resume text for agent analysis
        resume_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        
        # Run multi-agent analysis
        print("\nStep 2: Running multi-agent analysis...")
        self.agent_analysis = self.orchestrator.analyze_and_optimize(resume_text, job_description)
        
        # Extract keywords (for backward compatibility)
        keywords = self.extract_keywords(job_description)
        print(f"  ✓ Agent analysis complete")
        print(f"  ✓ ATS Score: {self.agent_analysis['ats_score']['overall_score']}/100")
        
        # Process
        print("\nStep 3: Analyzing and transforming content...")
        
        current_section = None
        summary_changed = False
        bullets_changed = 0
        title_changed = False
        self.used_phrases = []  # Track phrases used in this document to avoid repetition
        
        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            
            # DEBUG: Log ALL early paragraphs INCLUDING EMPTY to find title
            if idx < 10:
                if not text:
                    print(f"  [Para {idx}] EMPTY")
                else:
                    cap_words = len(re.findall(r'[A-Z][a-z]+', text))
                    non_alnum = len(re.findall(r'[^a-zA-Z0-9\s]', text))
                    print(f"  [Para {idx}] Len:{len(text)}, Cap:{cap_words}, NonAlnum:{non_alnum}")
                    print(f"           Text: '{text[:80]}'")
            
            if not text or len(text) < 3:
                continue
            
            # Check for role title: multiple capitalized phrases separated by non-alphanumeric chars
            # Pattern: "Word Word | Word Word | Word Word" or similar with any separator
            is_multipart_title = (
                idx < 5 and  # Early in document
                len(text) < 150 and  # Not too long
                len(re.findall(r'[A-Z][a-z]+', text)) >= 3 and  # Multiple capitalized words
                len(re.findall(r'[^a-zA-Z0-9\s]', text)) >= 1  # Has non-alphanumeric separators
            )
            
            # Title optimization disabled - not providing value
            # if not title_changed and is_multipart_title:
            #     Title optimization code removed
            
            # Check if section header
            is_header, section_type = self.is_section_header(text)
            if is_header:
                current_section = section_type
                print(f"  → Found {section_type.upper()} section at paragraph {idx}")
                continue
            
            # Process summary - use agent-generated summary
            if current_section == 'summary' and not summary_changed and len(text.strip()) > 30:
                # Get agent-generated summary
                new_summary = self.agent_analysis['new_summary'] if self.agent_analysis else self.create_new_summary(job_description, keywords, text)
                
                # DEBUG: Show what summary is being applied
                print(f"\n    📝 APPLYING SUMMARY:")
                print(f"       Original (first 100 chars): {text[:100]}...")
                print(f"       New (first 100 chars): {new_summary[:100]}...")
                
                if para.runs:
                    # Preserve formatting
                    font_name = para.runs[0].font.name
                    font_size = para.runs[0].font.size
                    font_bold = para.runs[0].font.bold
                    
                    para.clear()
                    run = para.add_run(new_summary)
                    
                    if font_name:
                        run.font.name = font_name
                    if font_size:
                        run.font.size = font_size
                    if font_bold:
                        run.font.bold = font_bold
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    
                    summary_changed = True
                    print(f"    ✓ Rewrote summary with agent intelligence ({len(new_summary.split())} words)")
            
            # Process experience content - enhance ALL content, not just bullets
            if current_section == 'experience':
                is_bullet = self.is_bullet_point(text)
                
                # Try to enhance ANY experience content if it has keywords
                if is_bullet or (len(text.split()) > 10 and len(text.split()) < 100):
                    # Remove bullet marker to get clean text
                    clean_text = re.sub(r'^[•\-\*\◦\·\►\▪]\s*', '', text.strip())
                    
                    # Use agent optimizer if available
                    if self.agent_analysis:
                        print(f"\n      Processing: {clean_text[:60]}...")
                        enhanced = self.orchestrator.optimize_bullet(
                            clean_text,
                            self.agent_analysis['job_analysis'],
                            self.agent_analysis['gap_analysis'],
                            self.agent_analysis.get('priority_keywords', [])
                        )
                    else:
                        enhanced = self.enhance_bullet(clean_text, keywords)
                    
                    if enhanced != clean_text and para.runs:
                        # Preserve formatting
                        font_name = para.runs[0].font.name
                        font_size = para.runs[0].font.size
                        font_bold = para.runs[0].font.bold
                        
                        para.clear()
                        # Restore bullet marker if it was a bullet
                        if is_bullet:
                            run = para.add_run(enhanced)
                        else:
                            run = para.add_run(enhanced)
                        
                        if font_name:
                            run.font.name = font_name
                        if font_size:
                            run.font.size = font_size
                        if font_bold:
                            run.font.bold = font_bold
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        
                        bullets_changed += 1
                        content_type = "bullet" if is_bullet else "paragraph"
                        if bullets_changed <= 5:
                            print(f"    ✓ Enhanced {content_type} {bullets_changed}")
        
        if bullets_changed > 5:
            print(f"    ✓ ... and {bullets_changed - 5} more bullets")
        
        # Save
        print(f"\nStep 4: Saving...")
        doc.save(output_path)
        print(f"  ✓ Saved to: {os.path.basename(output_path)}")
        
        # Results
        print("\n" + "="*70)
        print("RESULTS:")
        print("="*70)
        print(f"  ✓ Summary: {'Rewritten' if summary_changed else 'Not found'}")
        print(f"  ✓ Bullets enhanced: {bullets_changed}")
        print(f"  ✓ Keywords used: {len(keywords)}")
        print(f"  ✓ All changes in RED")
        print("="*70 + "\n")
        
        return {
            'success': True,
            'summary_changed': summary_changed,
            'bullets_changed': bullets_changed,
            'keywords_count': len(keywords),
            'output_path': output_path,
            'recommendations': [
                f"Summary: {'Rewritten' if summary_changed else 'Not changed'}",
                f"Bullets enhanced: {bullets_changed}",
                f"Keywords from job: {len(keywords)}",
                "All changes in RED",
                "Formatting preserved"
            ]
        }


def read_file_smart(filepath: str) -> str:
    """Read file content."""
    file_ext = os.path.splitext(filepath)[1].lower()
    if file_ext == '.docx':
        doc = Document(filepath)
        return '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
    else:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
